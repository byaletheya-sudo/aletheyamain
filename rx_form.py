"""Aletheya Toolbox · Medical / Psychiatric Update ("RX Update") generator.

The original form is an image-only scan, so we recreate it cleanly and render to
PDF (reportlab) or Word (python-docx) from one structured data dict.

DOCTORS is the provider directory extracted from the clinic's MD list — picking
one fills the Primary Physician section, the "Dear Dr." line, and the fax target.
"""
import io

# Provider directory. Picking one auto-fills the physician section.
DOCTORS = [
    {"first": "Sora", "last": "Lee", "prefix": "", "specialty": "", "company": "Lee Renal Care",
     "addr1": "18370 Burbank Blvd.", "addr2": "Suite 211", "city": "Tarzana", "state": "CA", "zip": "91356",
     "phone": "(747) 201-7444", "fax": "(818) 421-4256", "npi": "1891939237"},
    {"first": "Gemelia", "last": "Aguilera", "prefix": "", "specialty": "MD", "company": "",
     "addr1": "3325 Wilshire Blvd.", "addr2": "Suite 208", "city": "Los Angeles", "state": "CA", "zip": "90010",
     "phone": "(213) 465-2643", "fax": "(213) 232-4944", "npi": "1972616100"},
    {"first": "Kamran", "last": "Kamrava", "prefix": "Dr.", "specialty": "", "company": "",
     "addr1": "6915 Reseda Blvd", "addr2": "", "city": "Reseda", "state": "CA", "zip": "91335",
     "phone": "(818) 343-2121", "fax": "(818) 705-1622", "npi": "1053483800"},
    {"first": "Nicola", "last": "Azar", "prefix": "", "specialty": "MD", "company": "",
     "addr1": "1000 Newbury Rd Ste 265", "addr2": "", "city": "Thousand Oaks", "state": "CA", "zip": "91320",
     "phone": "(818) 914-4366", "fax": "(818) 914-4230", "npi": "1922259787"},
    {"first": "Peiman", "last": "Berdjis", "prefix": "", "specialty": "NP", "company": "",
     "addr1": "6222 Wilshire Blvd Suite 303", "addr2": "", "city": "Los Angeles", "state": "CA", "zip": "90048",
     "phone": "(323) 525-1999", "fax": "(323) 525-1991", "npi": "1326150590"},
    {"first": "Sameh", "last": "Shenouda", "prefix": "Dr.", "specialty": "Nephrologist", "company": "",
     "addr1": "7963 Van Nuys Blvd", "addr2": "#101", "city": "Panorama City", "state": "CA", "zip": "91402",
     "phone": "(818) 988-9818", "fax": "(818) 988-9828", "npi": "1487936902"},
    {"first": "Katherine", "last": "Guardado", "prefix": "", "specialty": "", "company": "",
     "addr1": "18250 Roscoe Blvd Ste 200", "addr2": "", "city": "Northridge", "state": "CA", "zip": "91325",
     "phone": "(818) 721-4800", "fax": "(818) 721-4825", "npi": "1821516535"},
]


def doctor_label(d):
    name = " ".join(x for x in [d.get("first", ""), d.get("last", "")] if x).strip()
    if d.get("specialty"):
        name += f", {d['specialty']}"
    where = d.get("city", "")
    return f"{name}" + (f" — {where}" if where else "")


def doctor_address(d):
    line1 = ", ".join(x for x in [d.get("addr1", ""), d.get("addr2", "")] if x)
    csz = f"{d.get('city','')}, {d.get('state','')} {d.get('zip','')}".strip().strip(",")
    return ", ".join(x for x in [line1, csz] if x).strip().strip(",")


def doctor_name(d):
    n = " ".join(x for x in [d.get("first", ""), d.get("last", "")] if x).strip()
    if d.get("specialty"):
        n += f", {d['specialty']}"
    return n


# Static clinic header info (from the form).
CLINIC = {
    "name": "Mountainview Adhc",
    "addr1": "23751 Roscoe Blvd",
    "addr2": "West Hills  CA  913043041",
    "phone": "Phone: (818) 312-0663",
    "fax": "Fax: (818) 716-8030",
    "fax_number": "(818) 716-8030",
    "rn": "Arvin Fekri, RN",
}

DISCLAIMER = (
    "Your patient is currently on one or more of the following treatments: Nursing, PT, OT, "
    "Maintenance, SW, ST, Psych, RD counseling and personal care. All Skilled Treatment care "
    "plans are designed to address mainly chronic issues. If any acute issues arise, we will "
    "ensure to refer them back to you for your immediate attention. In addition, if you wish to "
    "see your clients at your office on a regular basis, please call our Social Work Department "
    "and they will gladly assist you with your needs."
)


def _g(d, *keys, default=""):
    for k in keys:
        v = d.get(k)
        if v not in (None, ""):
            return v
    return default


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def render_rx_pdf(data):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    W, H = letter  # 612 x 792

    def L(x, y, t, font="Helvetica", size=10):
        c.setFont(font, size); c.drawString(x, y, str(t))

    def C(x, y, t, font="Helvetica", size=10):
        c.setFont(font, size); c.drawCentredString(x, y, str(t))

    def R(x, y, t, font="Helvetica", size=10):
        c.setFont(font, size); c.drawRightString(x, y, str(t))

    doc = data.get("doctor") or {}
    LM, RM = 54, 558  # margins

    # ---- Header ----
    C(W / 2, 752, CLINIC["name"], "Helvetica-Bold", 15)
    L(LM, 738, CLINIC["addr1"], "Times-Roman", 10)
    L(LM, 727, CLINIC["addr2"], "Times-Roman", 10)
    C(W / 2, 727, "Medical / Psychiatric Update", "Times-Bold", 13)
    R(RM, 740, CLINIC["phone"], "Times-Roman", 10)
    R(RM, 729, CLINIC["fax"], "Times-Roman", 10)

    # ---- Primary physician table ----
    L(LM, 710, "Primary Physician/Psychiatrist", "Times-Roman", 11)
    tx0, tx1 = 110, RM
    cN, cA = 250, 470          # column dividers
    ytop, hh, rh = 705, 15, 16
    # header row
    c.rect(tx0, ytop - hh, tx1 - tx0, hh, stroke=1, fill=0)
    C((tx0 + cN) / 2, ytop - hh + 4, "Name", "Times-Roman", 10)
    C((cN + cA) / 2, ytop - hh + 4, "Address", "Times-Roman", 10)
    C((cA + tx1) / 2, ytop - hh + 4, "Phone", "Times-Roman", 10)
    for dvx in (cN, cA):
        c.line(dvx, ytop - hh, dvx, ytop - hh - rh * 2)
    # two data rows
    for i in range(2):
        c.rect(tx0, ytop - hh - rh * (i + 1), tx1 - tx0, rh, stroke=1, fill=0)
    if doc:
        yb = ytop - hh - rh + 4
        L(tx0 + 4, yb, doctor_name(doc), "Times-Roman", 8.5)
        # address may be long; shrink to fit
        addr = doctor_address(doc)
        c.setFont("Times-Roman", 8.5)
        while c.stringWidth(addr, "Times-Roman", 8.5) > (cA - cN - 8) and len(addr) > 4:
            addr = addr[:-2]
        L(cN + 4, yb, addr, "Times-Roman", 8.5)
        L(cA + 4, yb, _g(doc, "phone"), "Times-Roman", 8.5)

    # ---- Dear Dr / assessment ----
    dear = "Dear Dr. " + (_g(doc, "last") or "")
    L(LM, 648, dear, "Times-Bold", 11)
    L(LM + 18, 630, "Your patient will be assessed by Mountainview ADHC multi-disciplinary team on "
      + (_g(data, "assess_date") or "________"), "Times-Roman", 10.5)

    L(LM + 18, 610, "Patient Name", "Times-Roman", 11)
    L(LM + 95, 610, _g(data, "patient_name"), "Times-Bold", 11)
    L(350, 610, "Date of Birth:", "Times-Roman", 11)
    L(425, 610, _g(data, "dob"), "Times-Bold", 11)

    L(LM + 18, 592, "Thank you so much for your assistance in updating this file!", "Times-Roman", 10.5)
    L(LM + 18, 579, "If we don't hear back from you within 2 weeks, we will be happy to consider all info as current",
      "Times-Roman", 10)

    # disclaimer (wrapped, italic)
    from reportlab.lib.utils import simpleSplit
    yy = 566
    for ln in simpleSplit(DISCLAIMER, "Times-Italic", 8.3, RM - (LM + 24)):
        L(LM + 24, yy, ln, "Times-Italic", 8.3); yy -= 10

    # ---- DISALLOW box ----
    by = yy - 4
    c.rect(150, by - 30, 320, 30, stroke=1, fill=0)
    C(310, by - 11, "Please Check if you DISALLOW Skilled PT or Skilled OT or Skilled ST", "Times-Italic", 8.5)

    def box(x, y, checked):
        c.rect(x, y, 9, 9, stroke=1, fill=0)
        if checked:
            c.setFont("Helvetica-Bold", 9); c.drawString(x + 1.2, y + 1, "X")

    box(168, by - 27, data.get("disallow_pt")); L(182, by - 26, "DISALLOW  PT", "Helvetica-Bold", 8.5)
    box(268, by - 27, data.get("disallow_ot")); L(282, by - 26, "DISALLOW  OT", "Helvetica-Bold", 8.5)
    box(372, by - 27, data.get("disallow_st")); L(386, by - 26, "DISALLOW  ST", "Helvetica-Bold", 8.5)

    # ---- Diagnoses ----
    dy = by - 44
    L(LM, dy, "Current Diagnosis: (per our records)  Please update using specific ICD codes:", "Helvetica", 10)
    diags = (data.get("diagnoses") or [])[:8]
    while len(diags) < 8:
        diags.append({"name": "", "icd": ""})
    gx0, gx1 = LM, RM
    icdx = RM - 90
    grow = 15
    gtop = dy - 6
    for i, dx in enumerate(diags):
        ry = gtop - grow * (i + 1)
        c.rect(gx0, ry, gx1 - gx0, grow, stroke=1, fill=0)
        c.line(icdx, ry, icdx, ry + grow)
        if dx.get("name"):
            L(gx0 + 4, ry + 4, dx["name"], "Times-Roman", 9)
        if dx.get("icd"):
            L(icdx + 4, ry + 4, dx["icd"], "Times-Roman", 9)

    # ---- Medications ----
    my = gtop - grow * 8 - 16
    L(LM, my, "Current Medications: (per our records)  Please update or verify:", "Helvetica", 10)
    meds = (data.get("medications") or [])[:11]
    while len(meds) < 11:
        meds.append("")
    mtop = my - 6
    for i, m in enumerate(meds):
        ry = mtop - grow * (i + 1)
        c.rect(gx0, ry, gx1 - gx0, grow, stroke=1, fill=0)
        if m:
            L(gx0 + 4, ry + 4, m, "Times-Roman", 9)

    # ---- Significant events ----
    sy = mtop - grow * 11 - 16
    L(LM + 6, sy, "Significant events over the last assessment period (6 months):", "Times-Roman", 9.5)
    ev = _g(data, "significant_events")
    if ev:
        for ln in simpleSplit(ev, "Times-Roman", 9.5, RM - LM - 12):
            sy -= 12; L(LM + 6, sy, ln, "Times-Roman", 9.5)
    c.line(LM + 6, sy - 6, RM, sy - 6)

    # ---- Sign / fax + dates ----
    fy = sy - 26
    L(LM + 18, fy, f"Please sign and fax this form to RN at Mountainview ADHC @ {CLINIC['fax_number']}. Thank you so",
      "Helvetica", 11)
    L(LM + 18, fy - 14, "much.", "Helvetica", 11)

    fax_text = f"This fax sent on behalf of {CLINIC['rn']}"
    L(LM, fy - 32, fax_text, "Times-Roman", 10)
    fw = c.stringWidth(fax_text, "Times-Roman", 10)
    c.line(LM, fy - 35, LM + fw, fy - 35)             # underline sits just BELOW the text

    the_date = _g(data, "date")
    # top date blank (to the right of the fax line)
    c.line(420, fy - 35, 545, fy - 35)
    if the_date:
        C(482, fy - 32, the_date, "Times-Roman", 10)
    C(482, fy - 47, "Date", "Times-Roman", 10)

    # MD signature + bottom date
    sig_label = "MD Signature:"
    L(LM + 40, fy - 66, sig_label, "Helvetica-Bold", 11)
    sw = c.stringWidth(sig_label, "Helvetica-Bold", 11)
    c.line(LM + 40 + sw + 8, fy - 66, 400, fy - 66)   # signature blank starts after the label
    c.line(420, fy - 66, 545, fy - 66)
    if the_date:
        C(482, fy - 63, the_date, "Times-Roman", 10)
    C(482, fy - 78, "Date", "Times-Roman", 10)

    c.showPage()
    c.save()
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def render_rx_docx(data):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = data.get("doctor") or {}
    document = Document()
    for s in document.sections:
        s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

    def para(text="", align=None, bold=False, italic=False, size=10, space_after=2):
        p = document.add_paragraph()
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        if text:
            r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        return p

    # Header — three-cell table for left / center / right
    htab = document.add_table(rows=1, cols=3)
    hc = htab.rows[0].cells
    p = hc[0].paragraphs[0]; r = p.add_run(CLINIC["addr1"] + "\n" + CLINIC["addr2"]); r.font.size = Pt(9)
    p = hc[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(CLINIC["name"]); r.bold = True; r.font.size = Pt(15)
    p2 = hc[1].add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Medical / Psychiatric Update"); r.bold = True; r.font.size = Pt(12)
    p = hc[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(CLINIC["phone"] + "\n" + CLINIC["fax"]); r.font.size = Pt(9)

    para("Primary Physician/Psychiatrist", size=10, space_after=2)
    ptab = document.add_table(rows=3, cols=3); ptab.style = "Table Grid"
    hdr = ptab.rows[0].cells
    for i, t in enumerate(["Name", "Address", "Phone"]):
        cp = hdr[i].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(t); run.bold = True; run.font.size = Pt(9)
    if doc:
        row = ptab.rows[1].cells
        for i, val in enumerate([doctor_name(doc), doctor_address(doc), _g(doc, "phone")]):
            cp = row[i].paragraphs[0]; run = cp.add_run(val); run.font.size = Pt(9)

    para()
    p = document.add_paragraph(); r = p.add_run("Dear Dr. " + (_g(doc, "last") or "")); r.bold = True; r.font.size = Pt(11)
    para("Your patient will be assessed by Mountainview ADHC multi-disciplinary team on "
         + (_g(data, "assess_date") or "________"), size=10.5)

    p = document.add_paragraph()
    r = p.add_run("Patient Name  "); r.font.size = Pt(11)
    r = p.add_run(_g(data, "patient_name")); r.bold = True; r.font.size = Pt(11)
    r = p.add_run("          Date of Birth:  "); r.font.size = Pt(11)
    r = p.add_run(_g(data, "dob")); r.bold = True; r.font.size = Pt(11)

    para("Thank you so much for your assistance in updating this file!", size=10.5)
    para("If we don't hear back from you within 2 weeks, we will be happy to consider all info as current", size=10)
    para(DISCLAIMER, italic=True, size=8.5, space_after=6)

    dis = "DISALLOW Skilled (check any):    "
    dis += ("☒" if data.get("disallow_pt") else "☐") + " PT     "
    dis += ("☒" if data.get("disallow_ot") else "☐") + " OT     "
    dis += ("☒" if data.get("disallow_st") else "☐") + " ST"
    para(dis, size=10, space_after=6)

    para("Current Diagnosis: (per our records)  Please update using specific ICD codes:", size=10)
    diags = [d for d in (data.get("diagnoses") or []) if d.get("name") or d.get("icd")]
    n = max(len(diags), 6)
    dtab = document.add_table(rows=n, cols=2); dtab.style = "Table Grid"
    dtab.columns[0].width = Inches(5.6); dtab.columns[1].width = Inches(1.4)
    for i in range(n):
        dx = diags[i] if i < len(diags) else {}
        c0 = dtab.rows[i].cells[0].paragraphs[0]; r = c0.add_run(dx.get("name", "")); r.font.size = Pt(9)
        c1 = dtab.rows[i].cells[1].paragraphs[0]; r = c1.add_run(dx.get("icd", "")); r.font.size = Pt(9)

    para()
    para("Current Medications: (per our records)  Please update or verify:", size=10)
    meds = [m for m in (data.get("medications") or []) if m]
    n = max(len(meds), 8)
    mtab = document.add_table(rows=n, cols=1); mtab.style = "Table Grid"
    for i in range(n):
        cp = mtab.rows[i].cells[0].paragraphs[0]
        r = cp.add_run(meds[i] if i < len(meds) else ""); r.font.size = Pt(9)

    para()
    p = document.add_paragraph(); r = p.add_run("Significant events over the last assessment period (6 months):")
    r.font.size = Pt(9.5)
    para(_g(data, "significant_events"), size=9.5, space_after=10)

    para(f"Please sign and fax this form to RN at Mountainview ADHC @ {CLINIC['fax_number']}. Thank you so much.",
         size=11)
    p = document.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"This fax sent on behalf of {CLINIC['rn']}"); r.font.size = Pt(10); r.underline = True

    the_date = _g(data, "date")
    p = document.add_paragraph()
    r = p.add_run("Date: "); r.font.size = Pt(10)
    r = p.add_run(the_date or "____________"); r.font.size = Pt(10)

    para(space_after=8)
    p = document.add_paragraph()
    r = p.add_run("MD Signature: ______________________________      Date: "); r.bold = True; r.font.size = Pt(11)
    r = p.add_run(the_date or "____________"); r.font.size = Pt(10)

    out = io.BytesIO(); document.save(out); return out.getvalue()
